"""
Génération du rapport Word (Partie 2)
Utilise python-docx. Architecture OOP avec surcharge de méthodes
"""

import os
import io
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "..", "assets")
os.makedirs(OUTPUT_DIR, exist_ok=True)


class WordReportGenerator:
    """
    Génère un rapport Word structuré à partir des données d'analyse.
    Méthodes surchargées : _add_heading existe en version H1, H2, H3.
    """

    AUTHOR = "DOTSU olympe - Python Avancé"
    ACCENT_COLOR = RGBColor(0x6C, 0x63, 0xFF)  # Violet #6C63FF

    def __init__(self):
        self._doc = Document()
        self._setup_styles()

    #  Styles 
    def _setup_styles(self) -> None:
        """Configure les styles globaux du document."""
        style = self._doc.styles["Normal"]
        style.font.name = "Arial"
        style.font.size = Pt(11)

        # Marges
        for section in self._doc.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(3)
            section.right_margin = Cm(2.5)

    # Surcharge de _add_heading (niveaux 1/2/3) ─
    def _add_heading(self, text: str, level: int = 1,
                     bold: bool = True, italic: bool = False,
                     color: RGBColor = None) -> None:
        """
        Ajoute un titre. Surcharge : comportement différent selon level.
        level=1 → H1 grand, centré
        level=2 → H2 normal
        level=3 → H3 petit italique
        """
        p = self._doc.add_paragraph()

        if level == 1:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(22)
            run.font.color.rgb = color or self.ACCENT_COLOR

        elif level == 2:
            run = p.add_run(text)
            run.bold = bold
            run.italic = italic
            run.font.size = Pt(15)
            run.font.color.rgb = color or self.ACCENT_COLOR
            # Ligne sous le titre
            self._add_horizontal_rule()

        elif level == 3:
            run = p.add_run(text)
            run.bold = bold
            run.italic = True
            run.font.size = Pt(12)
            run.font.color.rgb = color or RGBColor(0x88, 0x88, 0xAA)

    def _add_horizontal_rule(self) -> None:
        """Ligne de séparation sous un titre H2."""
        p = self._doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "6C63FF")
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _add_text(self, text: str, bold: bool = False, italic: bool = False,
                  size: int = 11, color: RGBColor = None) -> None:
        p = self._doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color

    def _add_stat_row(self, label: str, value: str) -> None:
        """Ajoute une ligne label : valeur en gras."""
        p = self._doc.add_paragraph()
        run_label = p.add_run(f"{label} : ")
        run_label.bold = True
        run_label.font.size = Pt(11)
        run_value = p.add_run(value)
        run_value.font.size = Pt(11)

    # Construction du rapport 
    def generate(
        self,
        metadata: dict,
        analysis: dict,
        chart_path: str,
        image_path: str | None,
        logo_path: str | None,
        output_path: str | None = None,
    ) -> str:
        """
        Génère le rapport complet et le sauvegarde.
        Retourne le chemin du fichier créé.
        """
        self._build_title_page(metadata, image_path)
        self._build_chart_page(analysis, chart_path)
        self._build_description_page(metadata, analysis)

        if output_path is None:
            output_path = os.path.join(OUTPUT_DIR, "rapport_peter_pan.docx")

        self._doc.save(output_path)
        return output_path

    #  Page de titre 
    def _build_title_page(self, metadata: dict, image_path: str | None) -> None:
        """Page 1 : titre du livre, image, auteur, auteur du rapport."""
        self._doc.add_paragraph()
        self._doc.add_paragraph()

        self._add_heading("Rapport d'Analyse Littéraire", level=1)
        self._doc.add_paragraph()

        # Titre du livre
        self._add_heading(metadata.get("title", "Peter Pan"), level=1,
                          color=RGBColor(0xFF, 0xC8, 0x57))

        # Image du livre
        if image_path and os.path.exists(image_path):
            p_img = self._doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try:
                run = p_img.add_run()
                run.add_picture(image_path, width=Inches(3.5))
            except Exception as e:
                print(f"Erreur image: {e}")
        self._doc.add_paragraph()

        # Auteur du livre
        p_author = self._doc.add_paragraph()
        p_author.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p_author.add_run(f"par {metadata.get('author', 'J. M. Barrie')}")
        r.bold = True
        r.font.size = Pt(14)
        r.font.color.rgb = RGBColor(0xCC, 0xCC, 0xFF)

        self._doc.add_paragraph()

        # Auteur du rapport
        p_rpt = self._doc.add_paragraph()
        p_rpt.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p_rpt.add_run(f"Rapport rédigé par : {self.AUTHOR}")
        r2.italic = True
        r2.font.size = Pt(11)
        r2.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

        # Date
        p_date = self._doc.add_paragraph()
        p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r3 = p_date.add_run(f"Date : {datetime.now().strftime('%d/%m/%Y')}")
        r3.italic = True
        r3.font.size = Pt(10)
        r3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        self._doc.add_page_break()

    # Page graphique 
    def _build_chart_page(self, analysis: dict, chart_path: str) -> None:
        """Page 2 : graphique de distribution + description."""
        self._add_heading("Distribution des longueurs de paragraphes", level=2)

        self._add_heading("Chapitre I – Peter Pan", level=3)
        self._doc.add_paragraph()

        # Graphique
        if chart_path and os.path.exists(chart_path):
            p_img = self._doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try:
                run = p_img.add_run()
                run.add_picture(chart_path, width=Inches(6))
            except Exception as e:
                print(f"Erreur image : {e}")
        self._doc.add_paragraph()

        self._doc.add_page_break()

    #  Page description 
    def _build_description_page(self, metadata: dict, analysis: dict) -> None:
        """Page 3 : analyse détaillée et statistiques."""
        self._add_heading("Description et Analyse", level=2)
        self._doc.add_paragraph()

        # Extrait du 1er chapitre de Peter Pan

        self._add_heading("Extrait du Chapitre I", level=3)
        paragraphs = analysis.get("paragraphs", [])
        if paragraphs:
            extrait = "\n\n".join(paragraphs[:3])
        else:
            extrait = "Texte non disponible."
        self._add_text(extrait)
        self._doc.add_paragraph()

        # Statistiques
        self._add_heading("Statistiques du Chapitre I", level=3)
        stats = [
            ("Nombre total de paragraphes",  str(analysis.get("total_paras", 0))),
            ("Nombre total de mots",         str(analysis.get("total_words", 0))),
            ("Mots minimum dans un §",       str(analysis.get("min_words", 0))),
            ("Mots maximum dans un §",       str(analysis.get("max_words", 0))),
            ("Mots en moyenne par §",        str(analysis.get("avg_words", 0))),
            ("Source",                       metadata.get("source", "Project Gutenberg")),
        ]
        for label, value in stats:
            self._add_stat_row(label, value)

        self._doc.add_paragraph()

        # Distribution
        self._add_heading("Distribution détaillée", level=3)
        dist = analysis.get("distribution", {})
        if dist:
            table = self._doc.add_table(rows=1, cols=2)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            hdr[0].text = "Nb de mots (arrondi)"
            hdr[1].text = "Nb de paragraphes"
            for nb_mots, nb_paras in sorted(dist.items()):
                row = table.add_row().cells
                row[0].text = str(nb_mots)
                row[1].text = str(nb_paras)
